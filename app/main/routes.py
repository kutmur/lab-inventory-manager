# app/main/routes.py

import io
from datetime import datetime
from flask import (
    render_template, redirect, url_for, flash, request, 
    send_file, current_app, stream_with_context, Response
)
from flask_login import login_required, current_user
from docx import Document
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from sqlalchemy.exc import SQLAlchemyError, IntegrityError, InvalidRequestError
import pytz
from sqlalchemy.orm import joinedload
from sqlalchemy.orm.exc import StaleDataError as ConcurrencyError

from app.main import bp
from app.main.forms import ProductForm, TransferForm, LabForm
from app.auth.decorators import admin_required
from app.models import Product, Lab, TransferLog, UserLog
from app.extensions import db, limiter
from app.utils import create_user_log
from app.socket_events import notify_inventory_update, notify_stock_alert


def format_timestamp(timestamp):
    """Convert UTC timestamp to Europe/Istanbul timezone.
    
    Args:
        timestamp: UTC datetime object
    
    Returns:
        datetime: Localized datetime in Europe/Istanbul timezone
    """
    istanbul_tz = pytz.timezone('Europe/Istanbul')
    return pytz.utc.localize(timestamp).astimezone(istanbul_tz)


def generate_excel(data):
    """Generate Excel file as a stream.
    
    Args:
        data: List of dictionaries containing product data
    
    Returns:
        BytesIO: Excel file stream
    """
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df = pd.DataFrame(data)
        df.to_excel(writer, sheet_name='Lab Inventory', index=False)
        workbook = writer.book
        worksheet = writer.sheets['Lab Inventory']
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#4F81BD',
            'font_color': 'white',
            'border': 1
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            max_len = max(
                df[value].astype(str).apply(len).max(),
                len(value)
            )
            worksheet.set_column(col_num, col_num, max_len + 2)
    
    output.seek(0)
    return output


def generate_pdf(data, lab_code=None):
    """Generate PDF file as a stream.
    
    Args:
        data: List of dictionaries containing product data
        lab_code: Optional lab code for title
    
    Returns:
        BytesIO: PDF file stream
    """
    buffer = io.BytesIO()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    elements = []
    styles = getSampleStyleSheet()
    title_text = (
        f"Lab Inventory Report - {lab_code}"
        if lab_code
        else "Full Inventory Report"
    )
    
    timestamp = format_timestamp(datetime.utcnow())
    elements.append(Paragraph(title_text, styles['Title']))
    elements.append(Paragraph(
        f"Generated on: {timestamp.strftime('%Y-%m-%d %H:%M')}",
        styles['Normal']
    ))
    
    # Create table data
    headers = [
        'Name', 'Registry #', 'Quantity', 'Unit',
        'Min Qty', 'Location', 'Notes'
    ]
    if not lab_code:
        headers.insert(0, 'Lab')
    
    table_data = [headers]
    for row in data:
        table_row = [
            row['Name'],
            row['Registry Number'],
            str(row['Quantity']),
            row['Unit'],
            str(row['Minimum Quantity']),
            row['Location'],
            row['Notes']
        ]
        if not lab_code:
            table_row.insert(0, row.get('Lab', ''))
        table_data.append(table_row)
    
    table = Table(table_data)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))
    elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    return buffer


def generate_word(data, lab_code=None):
    """Generate Word document as a stream.
    
    Args:
        data: List of dictionaries containing product data
        lab_code: Optional lab code for title
    
    Returns:
        BytesIO: Word document stream
    """
    doc = Document()
    title_text = (
        f"Lab Inventory Report - {lab_code}"
        if lab_code
        else "Full Inventory Report"
    )
    doc.add_heading(title_text, 0)
    
    timestamp = format_timestamp(datetime.utcnow())
    doc.add_paragraph(
        f"Generated on: {timestamp.strftime('%Y-%m-%d %H:%M')}"
    )
    
    headers = [
        'Name', 'Registry #', 'Quantity', 'Unit',
        'Min Qty', 'Location', 'Notes'
    ]
    if not lab_code:
        headers.insert(0, 'Lab')
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    for row in data:
        row_cells = table.add_row().cells
        col = 0
        if not lab_code:
            row_cells[col].text = row.get('Lab', '')
            col += 1
        row_cells[col].text = row['Name']
        row_cells[col + 1].text = row['Registry Number']
        row_cells[col + 2].text = str(row['Quantity'])
        row_cells[col + 3].text = row['Unit']
        row_cells[col + 4].text = str(row['Minimum Quantity'])
        row_cells[col + 5].text = row['Location']
        row_cells[col + 6].text = row['Notes']
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@bp.route('/')
@bp.route('/index')
def index():
    """Render the home page."""
    return render_template('main/index.html', title='Home')


@bp.route('/admin/dashboard')
@login_required
@admin_required
def admin_dashboard():
    """Render the admin dashboard."""
    return render_template('admin/dashboard.html')


@bp.route('/user/profile')
@login_required
def user_profile():
    """Render the user profile page."""
    return render_template('user/profile.html')


@bp.route('/protected-route')
@login_required
def protected_route():
    """Render a protected route for logged-in users."""
    return render_template('protected.html')


@bp.route('/admin-only')
@login_required
@admin_required
def admin_only():
    """Render a page accessible only to admin users."""
    return render_template('admin.html')


@bp.route('/products')
@login_required
def products():
    """Render a product list accessible to all users."""
    return render_template('products.html')


@bp.route('/admin/manage-users')
@login_required
@admin_required
def manage_users():
    """Render the user management page for admin users."""
    return render_template('admin/manage_users.html')


@bp.route('/dashboard')
@login_required
def dashboard():
    """Render the main dashboard with lab inventory."""
    # Safety net to ensure we always have labs
    from app.models import Lab
    if Lab.query.count() == 0:
        Lab.get_predefined_labs()
        
    labs = Lab.query.order_by(Lab.code).all()
    selected_lab_code = request.args.get('lab', 'all')
    
    if selected_lab_code != 'all':
        selected_lab = Lab.query.filter_by(code=selected_lab_code).first()
        if selected_lab:
            products_by_location = Product.get_sorted_products(
                selected_lab.id
            )
        else:
            flash('Invalid lab code selected', 'error')
            return redirect(url_for('main.dashboard'))
    else:
        products_by_location = []
        selected_lab = None

    return render_template(
        'main/dashboard.html',
        title='Dashboard',
        labs=labs,
        selected_lab=selected_lab,
        selected_lab_code=selected_lab_code,
        products_by_location=products_by_location
    )


@bp.route('/product/add', methods=['GET', 'POST'])
@login_required
@limiter.limit("20 per hour")
def add_product():
    """Add a new product to inventory."""
    selected_lab_code = request.args.get('lab')
    if not selected_lab_code:
        flash('Please select a lab first', 'warning')
        return redirect(url_for('main.dashboard'))
    
    selected_lab = Lab.query.filter_by(code=selected_lab_code).first()
    if not selected_lab:
        flash('Invalid lab selected', 'error')
        return redirect(url_for('main.dashboard'))
    
    form = ProductForm()
    form.lab_id.data = selected_lab.id  # Pre-select the lab
    
    if form.validate_on_submit():
        loc_parts = form.location.data.split('-')
        
        if loc_parts[0] == 'workspace':
            location_type = 'workspace'
            location_number = None
            location_position = None
        else:
            location_type = 'cabinet'
            location_number = loc_parts[1]
            location_position = loc_parts[2]

        try:
            product = Product(
                name=form.name.data,
                registry_number=form.registry_number.data,
                quantity=int(form.quantity.data),  # Ensure integer
                unit=form.unit.data,
                minimum_quantity=int(form.minimum_quantity.data),  # Ensure integer
                location_type=location_type,
                location_number=location_number,
                location_position=location_position,
                notes=form.notes.data,
                lab_id=selected_lab.id
            )
            db.session.add(product)
            
            # Create user log entry
            create_user_log(
                user=current_user,
                action_type='add',
                product=product,
                lab=selected_lab,
                quantity=product.quantity,
                notes=f"Initial addition: {product.quantity} {product.unit}"
            )
            
            db.session.commit()
            
            # Send notifications
            notify_inventory_update(product.id, 'add', {
                'name': product.name,
                'quantity': product.quantity,
                'unit': product.unit,
                'location': product.get_location_display()
            })
            
            if product.quantity <= product.minimum_quantity:
                notify_stock_alert(product.id, {
                    'name': product.name,
                    'quantity': product.quantity,
                    'minimum': product.minimum_quantity,
                    'unit': product.unit,
                    'lab': selected_lab.code
                })
            
            flash('Product added successfully', 'success')
            return redirect(url_for(
                'main.dashboard',
                lab=selected_lab_code
            ))
            
        except ValueError as ve:
            current_app.logger.error(
                f"Validation error while creating product: {str(ve)}"
            )
            flash(f'Validation error: {str(ve)}', 'error')
        except IntegrityError:
            db.session.rollback()
            flash(
                'A product with this registry number already exists '
                'in this lab',
                'error'
            )
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(
                f"Error while creating product: {str(e)}"
            )
            flash('An error occurred while adding the product', 'error')

    return render_template(
        'main/product_form.html',
        title='Add Product',
        form=form,
        selected_lab=selected_lab,
        labs=[selected_lab]  # Pass labs for location dropdown JS
    )


@bp.route('/product/<int:id>/edit', methods=['GET', 'POST'])
@login_required
@limiter.limit("20 per hour")
def edit_product(id):
    """Edit an existing product."""
    product = Product.query.get_or_404(id)
    
    if not current_user.is_editor():
        flash('You do not have permission to edit products', 'error')
        return redirect(url_for('main.dashboard'))
    
    form = ProductForm(obj=product)
    
    # Pre-populate the lab_id field
    form.lab_id.data = product.lab_id
    
    if form.validate_on_submit():
        try:
            old_quantity = product.quantity
            loc_parts = form.location.data.split('-')
            
            # Direct attribute assignment instead of using update_record()
            product.name = form.name.data
            product.registry_number = form.registry_number.data
            product.quantity = int(form.quantity.data)  # Ensure integer
            product.unit = form.unit.data
            product.minimum_quantity = int(form.minimum_quantity.data)  # Ensure integer
            product.notes = form.notes.data
            
            if loc_parts[0] == 'workspace':
                product.location_type = 'workspace'
                product.location_number = None
                product.location_position = None
            else:
                product.location_type = 'cabinet'
                product.location_number = loc_parts[1]
                product.location_position = loc_parts[2]
            
            # Create user log for quantity change
            if product.quantity != old_quantity:
                quantity_change = product.quantity - old_quantity
                create_user_log(
                    user=current_user,
                    action_type='edit',
                    product=product,
                    lab=product.lab,
                    quantity=quantity_change,
                    notes=(
                        f"Quantity changed from {old_quantity} "
                        f"to {product.quantity}"
                    )
                )
            
            db.session.commit()
            
            # Send notifications
            notify_inventory_update(product.id, 'edit', {
                'name': product.name,
                'quantity': product.quantity,
                'unit': product.unit,
                'location': product.get_location_display()
            })
            
            if product.quantity <= product.minimum_quantity:
                notify_stock_alert(product.id, {
                    'name': product.name,
                    'quantity': product.quantity,
                    'minimum': product.minimum_quantity,
                    'unit': product.unit,
                    'lab': product.lab.code
                })
            
            flash('Product updated successfully', 'success')
            return redirect(url_for(
                'main.dashboard',
                lab=product.lab.code
            ))
            
        except ValueError as ve:
            db.session.rollback()
            current_app.logger.exception(
                f"Validation error while updating product: {str(ve)}"
            )
            flash(f'Validation error: {str(ve)}', 'error')
        except IntegrityError:
            db.session.rollback()
            flash('Duplicate registry number in this lab.', 'error')
        except SQLAlchemyError as e:
            db.session.rollback()
            current_app.logger.exception(e)
            flash('DB error while updating product.', 'error')
        except Exception as e:
            db.session.rollback()
            current_app.logger.exception(
                f"Error while updating product: {str(e)}"
            )
            flash('An error occurred while updating the product', 'error')
    
    # Pre-populate location field
    if not form.is_submitted():
        if product.location_type == 'workspace':
            form.location.data = 'workspace'
        else:
            form.location.data = (
                f"cabinet-{product.location_number}-"
                f"{product.location_position}"
            )

    return render_template(
        'main/product_form.html',
        title='Edit Product',
        form=form,
        selected_lab=product.lab,
        labs=[product.lab]  # Pass labs for location dropdown JS
    )


@bp.route('/product/<int:id>/delete', methods=['POST'])
@login_required
@admin_required
@limiter.limit("10 per hour")
def delete_product(id):
    """Delete a product from inventory."""
    product = Product.query.get_or_404(id)
    lab_code = product.lab.code
    
    # Store product information before deletion for logging
    product_name = product.name
    product_registry = product.registry_number
    product_quantity = product.quantity
    lab = product.lab
    
    try:
        # Log the deletion before actually deleting the product
        create_user_log(current_user, 'delete', product, lab,
                      -product_quantity, f"Product {product_name} (#{product_registry}) deleted from {lab_code}")
                      
        # Now delete the product
        db.session.delete(product)
        db.session.commit()
        flash('Product deleted successfully', 'success')
    except IntegrityError:
        db.session.rollback()
        flash('Cannot delete: product is referenced elsewhere.', 'error')
    except SQLAlchemyError as e:
        db.session.rollback()
        current_app.logger.exception(e)
        flash('DB error while deleting product.', 'error')
    
    return redirect(url_for('main.dashboard', lab=lab_code))


@bp.route('/product/<int:product_id>/transfer', methods=['GET', 'POST'])
@login_required
@limiter.limit("20 per hour")
def transfer_product(product_id):
    """Transfer a product between labs."""
    source_product = Product.query.join(Lab).filter(Product.id == product_id).first_or_404()
    
    if not source_product.lab:
        flash('Error: Source product has no associated laboratory.', 'danger')
        return redirect(url_for('main.dashboard'))
    
    destination_labs = Lab.query.filter(Lab.id != source_product.lab_id).all()
    if not destination_labs:
        flash('No available destination laboratories for transfer.', 'danger')
        return redirect(url_for('main.dashboard'))
    
    form = TransferForm(request.form, 
                       source_lab_id=source_product.lab_id, 
                       max_quantity=source_product.quantity,
                       product=source_product)

    if form.validate_on_submit():
        target_lab_id = form.destination_lab_id.data
        transfer_quantity = int(form.quantity.data)  # Ensure integer
        notes = form.notes.data

        try:
            destination_lab = Lab.query.get_or_404(target_lab_id)

            # Basic validation checks
            if transfer_quantity > source_product.quantity:
                flash('Transfer quantity cannot exceed available quantity!', 'danger')
                return render_template('main/transfer_form.html', title='Transfer Product', form=form, product=source_product)
                
            if transfer_quantity <= 0:
                flash('Transfer quantity must be positive!', 'danger')
                return render_template('main/transfer_form.html', title='Transfer Product', form=form, product=source_product)

            clean_registry = source_product.registry_number.strip()

            # ➡️ 1) Kaynak miktarı küçült
            if source_product.quantity < transfer_quantity:
                raise ValueError("Transfer quantity exceeds available stock")
            source_product.quantity -= transfer_quantity

            # ➡️ 2) Hedef ürünü bul / oluştur
            target_product = Product.query.filter_by(
                lab_id=target_lab_id,
                registry_number=clean_registry
            ).first()
            if target_product:
                target_product.quantity += transfer_quantity
            else:
                target_product = Product(
                    name=source_product.name,
                    registry_number=clean_registry,
                    quantity=transfer_quantity,
                    unit=source_product.unit,
                    minimum_quantity=source_product.minimum_quantity,
                    location_type='workspace',
                    notes=notes or source_product.notes,
                    lab_id=target_lab_id
                )
                db.session.add(target_product)

            # ➡️ 3) TransferLog + UserLog
            transfer_log = TransferLog(
                product_id           = source_product.id,
                source_lab_id        = source_product.lab_id,
                destination_lab_id   = target_lab_id,
                quantity             = transfer_quantity,
                notes                = f"{source_product.lab.code} ➜ {destination_lab.code}",
                created_by_id        = current_user.id
            )
            db.session.add(transfer_log)

            create_user_log(current_user, 'transfer', source_product,
                            source_product.lab, -transfer_quantity,
                            f"Sent to {destination_lab.code}")
            create_user_log(current_user, 'transfer', target_product,
                            destination_lab, transfer_quantity,
                            f"Received from {source_product.lab.code}")

            # 🔑 Tek seferde kaydet
            db.session.commit()

            notify_inventory_update(source_product.id, 'transfer', {
                'name'           : source_product.name,
                'quantity'       : source_product.quantity,
                'source_lab'     : source_product.lab_id,
                'destination_lab': target_lab_id
            })
            flash('Product transferred successfully!', 'success')
            return redirect(url_for('main.dashboard', lab=source_product.lab.code))

        except (InvalidRequestError, SQLAlchemyError) as db_err:
            db.session.rollback()
            current_app.logger.exception(f"DB error during transfer: {db_err}")
            flash('Database error during transfer. Please try again.', 'danger')

        except ValueError as ve:
            db.session.rollback()
            flash(str(ve), 'warning')

        except Exception as e:
            db.session.rollback()
            current_app.logger.exception(f"Unexpected error: {e}")
            flash('Unexpected error during transfer.', 'danger')

    return render_template('main/transfer_form.html', title='Transfer Product', form=form, product=source_product)


#######################################################################
#  - - -  BU İKİ ROUTE LAB-ID BAZLI DÜZENLEME/SİLME İSTEYENLER İÇİN - -
#######################################################################

@bp.route('/lab/<int:lab_id>/product/<int:product_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_lab_product(lab_id, product_id):
    """Edit a product in a specific lab."""
    lab = Lab.query.get_or_404(lab_id)
    product = Product.query.filter_by(id=product_id, lab_id=lab_id).first_or_404()

    form = ProductForm(obj=product)
    form.lab_id.choices = [(lab.id, f"{lab.code} - {lab.description}")]
    form.lab_id.data = lab.id

    if product.location_type == 'workspace':
        form.location_number.data = 'workspace'
    else:
        form.location_number.data = f"cabinet-{product.location_number}-{product.location_position}"

    if form.validate_on_submit():
        try:
            old_quantity = product.quantity

            loc_parts = form.location_number.data.split('-')
            if loc_parts[0] == 'workspace':
                product.location_type = 'workspace'
                product.location_number = None
                product.location_position = None
            else:
                product.location_type = 'cabinet'
                product.location_number = loc_parts[1]
                product.location_position = loc_parts[2]

            product.name = form.name.data
            product.registry_number = form.registry_number.data
            product.quantity = form.quantity.data
            product.unit = form.unit.data
            product.minimum_quantity = form.minimum_quantity.data
            product.notes = form.notes.data

            if old_quantity != product.quantity:
                diff = product.quantity - old_quantity
                create_user_log(
                    user=current_user,
                    action_type='edit',
                    product=product,
                    lab=lab,
                    quantity=diff,
                    notes="Edited product quantity"
                )

            db.session.commit()
            flash(f'Product in {lab.code} updated successfully!', 'success')
            return redirect(url_for('main.dashboard', lab=lab.code))

        except Exception as e:
            db.session.rollback()
            flash('Error updating product. Please try again.', 'error')

    return render_template('main/product_form.html', form=form,
                           title=f"Edit Product in {lab.code}",
                           labs=[lab])


@bp.route('/lab/<int:lab_id>/product/<int:product_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_lab_product(lab_id, product_id):
    """Delete a product from a specific lab."""
    try:
        lab = Lab.query.get_or_404(lab_id)
        product = Product.query.filter_by(id=product_id, lab_id=lab_id).first_or_404()
        
        # Store product information before deletion for logging
        product_name = product.name
        product_registry = product.registry_number
        product_quantity = product.quantity

        # Log the deletion before actually deleting the product
        create_user_log(
            user=current_user,
            action_type='delete',
            product=product,
            lab=lab,
            quantity=-product.quantity,
            notes=f"Deleted product {product_name} (#{product_registry}) from lab {lab.code}"
        )

        db.session.delete(product)
        db.session.commit()
        flash(f'Product "{product_name}" deleted successfully from {lab.code}!', 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.exception(e)
        flash('Error deleting product. Please try again.', 'error')

    return redirect(url_for('main.dashboard', lab=lab.code))


#######################################################################
#  TRANSFER ROUTE
#######################################################################

@bp.route('/transfer', methods=['GET', 'POST'])
@login_required
def transfer_between_labs():
    """Transfer products between labs."""
    form = TransferForm()
    form.source_lab_id.choices = [(lab.id, lab.name) for lab in Lab.query.all()]
    form.destination_lab_id.choices = [(lab.id, lab.name) for lab in Lab.query.all()]
    form.product_id.choices = [(p.id, f"{p.name} ({p.registry_number})") for p in Product.query.all()]

    if form.validate_on_submit():
        source_lab_id = form.source_lab_id.data
        destination_lab_id = form.destination_lab_id.data

        source_lab = Lab.query.get(source_lab_id)
        destination_lab = Lab.query.get(destination_lab_id)
        
        if not source_lab or not destination_lab:
            flash('Invalid source or destination laboratory.', 'error')
            return redirect(url_for('main.transfer_between_labs'))

        if source_lab_id == destination_lab_id:
            flash('Source and destination labs cannot be the same.', 'error')
            return redirect(url_for('main.transfer_between_labs'))

        try:
            db.session.begin_nested()

            product = Product.query.get_or_404(form.product_id.data)
            if not product or product.lab_id != source_lab_id:
                flash('Selected product does not belong to the source lab.', 'error')
                db.session.rollback()
                return redirect(url_for('main.transfer_between_labs'))

            transfer_qty = form.quantity.data
            if product.quantity < transfer_qty:
                flash('Insufficient quantity in the source lab.', 'error')
                db.session.rollback()
                return redirect(url_for('main.transfer_between_labs'))

            old_quantity = product.quantity
            product.quantity = old_quantity - transfer_qty
            db.session.flush()

            dest_product = Product.query.filter_by(registry_number=product.registry_number, lab_id=destination_lab_id).first()
            if dest_product:
                dest_product.quantity += transfer_qty
            else:
                dest_product = Product(
                    name=product.name,
                    registry_number=product.registry_number,
                    quantity=transfer_qty,
                    unit=product.unit,
                    minimum_quantity=product.minimum_quantity,
                    location_type='workspace',
                    location_number=None,
                    location_position=None,
                    notes=product.notes,
                    lab_id=destination_lab_id
                )
                db.session.add(dest_product)

            transfer_log = TransferLog(
                product_id=product.id,
                source_lab_id=source_lab_id,
                destination_lab_id=destination_lab_id,
                quantity=transfer_qty,
                notes=f"Transferred from {source_lab.code} to {destination_lab.code}",
                created_by_id=current_user.id
            )
            db.session.add(transfer_log)

            create_user_log(
                user=current_user,
                action_type='transfer',
                product=product,
                lab=source_lab,
                quantity=-transfer_qty,
                notes=f"Transferred out {transfer_qty} from {source_lab.code} to {destination_lab.code}"
            )
            create_user_log(
                user=current_user,
                action_type='transfer',
                product=dest_product,
                lab=destination_lab,
                quantity=transfer_qty,
                notes=f"Received {transfer_qty} from {source_lab.code}"
            )

            db.session.commit()

            notify_inventory_update(product.id, 'transfer', {
                'name': product.name,
                'quantity': product.quantity,
                'source_lab': source_lab_id,
                'destination_lab': destination_lab_id
            })

            flash('Transfer completed successfully!', 'success')
            return redirect(url_for('main.dashboard'))

        except ConcurrencyError:
            db.session.rollback()
            flash('Transfer failed: Product was modified by another user. Please try again.', 'error')
        except Exception as e:
            db.session.rollback()
            flash('Error processing transfer. Please try again.', 'error')

    return render_template('main/transfer_form.html', form=form, title='Transfer Product')


#######################################################################
#  - - - - - - -  EXPORT ROTALARI (PDF / XLSX / DOCX) - - - - - - - - -
#######################################################################

@bp.route('/export/<lab_code>/<format>')
@login_required
@limiter.limit("10 per minute")
def export_lab(lab_code, format):
    """Export lab inventory with streaming response."""
    lab = Lab.query.filter_by(code=lab_code).first_or_404()
    products = Product.query.filter_by(lab_id=lab.id)\
        .options(joinedload(Product.lab))\
        .all()

    data = []
    for product in products:
        data.append({
            'Name': product.name,
            'Registry Number': product.registry_number,
            'Quantity': product.quantity,
            'Unit': product.unit,
            'Minimum Quantity': product.minimum_quantity,
            'Location': product.get_location_display(),
            'Notes': product.notes or ''
        })

    timestamp = format_timestamp(datetime.utcnow())\
        .strftime('%Y%m%d_%H%M%S')
    filename = f"inventory_{lab.code}_{timestamp}"

    if format == 'xlsx':
        return Response(
            stream_with_context(generate_excel(data)),
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".spreadsheetml.sheet"
            ),
            headers={
                "Content-Disposition": f"attachment; filename={filename}.xlsx"
            }
        )
    elif format == 'pdf':
        return Response(
            stream_with_context(generate_pdf(data, lab.code)),
            mimetype='application/pdf',
            headers={
                "Content-Disposition": f"attachment; filename={filename}.pdf"
            }
        )
    elif format == 'docx':
        return Response(
            stream_with_context(generate_word(data, lab.code)),
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": f"attachment; filename={filename}.docx"
            }
        )

    return "Format not supported", 400


@bp.route('/export/all/<format>')
@login_required
@limiter.limit("5 per minute")
def export_all_labs(format):
    """Export all labs inventory with streaming response."""
    labs = Lab.query.all()
    all_data = []

    for lab in labs:
        products = Product.query.filter_by(lab_id=lab.id)\
            .options(joinedload(Product.lab))\
            .all()
            
        for product in products:
            all_data.append({
                'Lab': f"{lab.code} - {lab.description}",
                'Name': product.name,
                'Registry Number': product.registry_number,
                'Quantity': product.quantity,
                'Unit': product.unit,
                'Minimum Quantity': product.minimum_quantity,
                'Location': product.get_location_display(),
                'Notes': product.notes or ''
            })

    if not all_data:
        flash('No data available to export', 'warning')
        return redirect(url_for('main.dashboard'))

    timestamp = format_timestamp(datetime.utcnow())\
        .strftime('%Y%m%d_%H%M%S')
    filename = f"full_inventory_{timestamp}"

    if format == 'xlsx':
        return Response(
            stream_with_context(generate_excel(all_data)),
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".spreadsheetml.sheet"
            ),
            headers={
                "Content-Disposition": f"attachment; filename={filename}.xlsx"
            }
        )
    elif format == 'pdf':
        return Response(
            stream_with_context(generate_pdf(all_data)),
            mimetype='application/pdf',
            headers={
                "Content-Disposition": f"attachment; filename={filename}.pdf"
            }
        )
    elif format == 'docx':
        return Response(
            stream_with_context(generate_word(all_data)),
            mimetype=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            headers={
                "Content-Disposition": f"attachment; filename={filename}.docx"
            }
        )

    return "Format not supported", 400


#######################################################################
# USER ACTIVITY LOG GÖRÜNTÜLEME (ADMIN)
#######################################################################

@bp.route('/logs')
@login_required
@admin_required
def user_logs():
    """List user activity logs."""
    page = request.args.get('page', 1, type=int)
    logs = UserLog.query.order_by(
        UserLog.timestamp.desc()
    ).paginate(
        page=page,
        per_page=50
    )
    return render_template(
        'main/user_logs.html',
        title='User Activity Logs',
        logs=logs
    )


@bp.route('/search')
@login_required
@limiter.limit("30 per minute")
def search_products():
    """Search products by name or registry number."""
    query = request.args.get('q', '')
    lab_code = request.args.get('lab', 'all')
    
    if lab_code != 'all':
        lab = Lab.query.filter_by(code=lab_code).first_or_404()
        products = Product.search(query, lab.id)
    else:
        products = Product.search(query)
    
    return render_template('main/search_results.html',
                         title='Search Results',
                         query=query,
                         products=products,
                         selected_lab_code=lab_code)


# Commented out as per requirements to disable lab creation
# @bp.route('/lab/add', methods=['GET', 'POST'])
# @login_required
# @admin_required
# def add_lab():
#     """Add a new laboratory."""
#     form = LabForm()
#     if form.validate_on_submit():
#         try:
#             lab = Lab(
#                 name=form.name.data,
#                 description=form.description.data,
#                 location=form.location.data
#             )
#             db.session.add(lab)
#             db.session.commit()
#             flash('Laboratory added successfully', 'success')
#             return redirect(url_for('main.dashboard'))
#         except IntegrityError:
#             db.session.rollback()
#             flash('A lab with this name already exists', 'error')
#         except Exception as e:
#             db.session.rollback()
#             current_app.logger.error(f"Error adding lab: {str(e)}")
#             flash('An error occurred while adding the lab', 'error')
#     
#     return render_template(
#         'main/lab_form.html',
#         title='Add Laboratory',
#         form=form
#     )
