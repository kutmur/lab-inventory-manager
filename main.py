import tkinter as tk
from tkinter import ttk
from tkinter import *
from tkinter import messagebox
from tkinter import simpledialog
import json
import os
from datetime import datetime
import pandas as pd 
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

class EditDialog:
    def __init__(self, parent, title, current_name, current_quantity):
        self.dialog = tk.Toplevel(parent)
        self.dialog.title(title)
        self.dialog.geometry("300x200")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        # Center the dialog
        self.dialog.geometry("+%d+%d" % (
            parent.winfo_rootx() + parent.winfo_width()/2 - 150,
            parent.winfo_rooty() + parent.winfo_height()/2 - 100
        ))
        
        # Product Name
        ttk.Label(self.dialog, text="Product Name:").pack(pady=(10, 0))
        self.name_var = tk.StringVar(value=current_name)
        self.name_entry = ttk.Entry(self.dialog, textvariable=self.name_var, width=30)
        self.name_entry.pack(pady=(0, 10))
        
        # Quantity
        ttk.Label(self.dialog, text="Quantity:").pack(pady=(10, 0))
        self.quantity_var = tk.StringVar(value=str(current_quantity))
        self.quantity_entry = ttk.Entry(self.dialog, textvariable=self.quantity_var, width=30)
        self.quantity_entry.pack(pady=(0, 10))
        
        # Buttons
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text="Save", command=self.save).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Cancel", command=self.cancel).pack(side=LEFT, padx=5)
        
        self.result = None
        
    def save(self):
        try:
            quantity = int(self.quantity_var.get())
            if quantity <= 0:
                messagebox.showerror("Error", "Quantity must be a positive number")
                return
            
            name = self.name_var.get().strip()
            if not name:
                messagebox.showerror("Error", "Product name cannot be empty")
                return
            
            self.result = {
                "name": name,
                "quantity": quantity
            }
            self.dialog.destroy()
            
        except ValueError:
            messagebox.showerror("Error", "Please enter a valid number for quantity")
    
    def cancel(self):
        self.dialog.destroy()

class LoginWindow:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Lab Inventory Manager - Login")
        self.root.geometry("300x200")
        
        # Center the window
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - 300) // 2
        y = (screen_height - 200) // 2
        self.root.geometry(f"300x200+{x}+{y}")
        
        # Create main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=BOTH, expand=True)
        
        # Username
        ttk.Label(main_frame, text="Username:").pack(pady=(0, 5))
        self.username_var = tk.StringVar()
        self.username_entry = ttk.Entry(main_frame, textvariable=self.username_var, width=30)
        self.username_entry.pack(pady=(0, 10))
        
        # Password
        ttk.Label(main_frame, text="Password:").pack(pady=(0, 5))
        self.password_var = tk.StringVar()
        self.password_entry = ttk.Entry(main_frame, textvariable=self.password_var, show="*", width=30)
        self.password_entry.pack(pady=(0, 20))
        
        # Login button
        self.login_button = ttk.Button(main_frame, text="Login", command=self.validate_login)
        self.login_button.pack()
        
        # Bind Enter key to login
        self.root.bind('<Return>', lambda e: self.validate_login())
        
        # Focus username entry
        self.username_entry.focus()
        
        # Authentication result
        self.authenticated = False
        self.user_info = None
        
        # Load users file
        self.users_file = 'users.json'
        self.users = self.load_users()
    
    def load_users(self):
        """Load users from JSON file or create default if not exists"""
        default_users = {
            "admin": {"password": "1234", "role": "admin"},
            "teacher": {"password": "abcd", "role": "user"}
        }
        
        try:
            if os.path.exists(self.users_file):
                with open(self.users_file, 'r') as f:
                    return json.load(f)
            else:
                # Create default users file
                with open(self.users_file, 'w') as f:
                    json.dump(default_users, f, indent=4)
                return default_users
        except Exception as e:
            print(f"Error loading users file: {e}")
            return default_users
        
    def validate_login(self):
        username = self.username_var.get().strip()
        password = self.password_var.get().strip()
        
        if not username or not password:
            messagebox.showerror("Login Failed", "Please enter both username and password")
            return
        
        if username not in self.users:
            messagebox.showerror("Login Failed", "User not found")
            self.password_var.set("")
            return
        
        user_data = self.users[username]
        if password != user_data["password"]:
            messagebox.showerror("Login Failed", "Invalid password")
            self.password_var.set("")
            return
        
        self.authenticated = True
        self.user_info = {
            "username": username,
            "role": user_data["role"]
        }
        self.root.destroy()
    
    def run(self):
        self.root.mainloop()
        return self.authenticated, self.user_info

class LogViewer:
    def __init__(self, parent):
        # Create Toplevel window
        self.window = tk.Toplevel(parent)
        self.window.title("User Action Logs")
        self.window.geometry("800x500")
        
        # Center the window relative to parent
        self.window.geometry("+%d+%d" % (
            parent.winfo_rootx() + parent.winfo_width()/2 - 400,
            parent.winfo_rooty() + parent.winfo_height()/2 - 250
        ))
        
        # Create main frame
        main_frame = ttk.Frame(self.window, padding="10")
        main_frame.pack(fill=BOTH, expand=True)
        
        # Create Treeview with scrollbar
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill=BOTH, expand=True)
        
        # Create Treeview
        self.tree = ttk.Treeview(
            tree_frame,
            columns=("User", "Action", "Registry", "Lab", "Quantity", "Timestamp"),
            show="headings"
        )
        
        # Configure column headings
        for col in self.tree["columns"]:
            self.tree.heading(col, text=col)
            # Adjust column widths
            if col in ["User", "Action", "Registry", "Lab", "Quantity"]:
                self.tree.column(col, width=100)
            else:  # Timestamp
                self.tree.column(col, width=150)
        
        # Add scrollbar
        scrollbar = ttk.Scrollbar(tree_frame, orient=VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # Pack Treeview and scrollbar
        self.tree.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        # Load and display logs
        self.load_logs()
        
        # Add close button at bottom
        close_button = ttk.Button(main_frame, text="Close", command=self.window.destroy)
        close_button.pack(pady=(10, 0))
        
        # Make window modal
        self.window.transient(parent)
        self.window.grab_set()
    
    def load_logs(self):
        """Load and display logs from user_log.json"""
        log_file = 'user_log.json'
        
        if not os.path.exists(log_file):
            self.show_no_logs_message()
            return
        
        try:
            with open(log_file, 'r') as f:
                logs = json.load(f)
            
            if not logs:
                self.show_no_logs_message()
                return
            
            # Insert logs into Treeview
            for log in reversed(logs):  # Show newest first
                self.tree.insert("", "end", values=(
                    log["user"],
                    log["action"],
                    log["registry"],
                    log["lab"],
                    log["quantity"],
                    log["timestamp"]
                ))
                
        except Exception as e:
            print(f"Error loading logs: {e}")
            self.show_no_logs_message()
    
    def show_no_logs_message(self):
        """Display message when no logs are available"""
        # Remove all columns
        for col in self.tree["columns"]:
            self.tree.heading(col, text="")
            self.tree.column(col, width=0)
        
        # Add single column for message
        self.tree["columns"] = ("message",)
        self.tree.column("message", width=400)
        self.tree.heading("message", text="")
        
        # Insert message
        self.tree.insert("", "end", values=("No logs available",))

class LabInventoryManager:
    def __init__(self, root, user_info):
        self.root = root
        self.root.title("Lab Inventory Management System")
        self.root.geometry("1000x600")
        
        # Store user info
        self.user_info = user_info
        
        # Initialize log file
        self.user_log_file = 'user_log.json'
        self.user_log = self.load_user_log()
        
        # Create user info and view logs frame
        top_frame = ttk.Frame(self.root)
        top_frame.pack(side=TOP, fill=X, padx=10, pady=5)
        
        # Add View Logs button (only for admin)
        if self.user_info["role"] == "admin":
            self.view_logs_button = ttk.Button(top_frame, text="View Logs", command=self.show_logs)
            self.view_logs_button.pack(side=RIGHT, padx=(10, 0))
        
        # User info label
        user_label = ttk.Label(
            top_frame,
            text=f"Logged in as: {self.user_info['username']} ({self.user_info['role']})",
            font=('Helvetica', 10)
        )
        user_label.pack(side=RIGHT)
        
        # Initialize inventory dictionary and transfer log
        self.inventory_file = 'inventory.json'
        self.transfer_log_file = 'transfer_log.json'
        self.inventory = self.load_inventory()
        self.transfer_log = self.load_transfer_log()
        
        # Initialize export directories
        self.export_base_dir = 'exports'
        self.export_dirs = {
            'pdf': os.path.join(self.export_base_dir, 'pdf'),
            'word': os.path.join(self.export_base_dir, 'word'),
            'excel': os.path.join(self.export_base_dir, 'excel')
        }
        self.create_export_directories()
        
        # Left Frame - Input Section
        left_frame = ttk.Frame(root, padding="10")
        left_frame.pack(side=LEFT, fill=Y, padx=10, pady=10)
        
        # Lab Selection
        ttk.Label(left_frame, text="Select Lab:").pack(anchor=W)
        self.lab_combo = ttk.Combobox(left_frame, width=25, values=list(self.inventory.keys()))
        self.lab_combo.pack(pady=(0, 10), anchor=W)
        self.lab_combo.set("Select Lab")  # Default text
        self.lab_combo.bind('<<ComboboxSelected>>', self.on_lab_selected)
        
        # Product Details
        ttk.Label(left_frame, text="Product Name:").pack(anchor=W)
        self.product_entry = ttk.Entry(left_frame, width=30)
        self.product_entry.pack(pady=(0, 10), anchor=W)
        
        ttk.Label(left_frame, text="Registry Number:").pack(anchor=W)
        self.registry_entry = ttk.Entry(left_frame, width=30)
        self.registry_entry.pack(pady=(0, 10), anchor=W)
        
        ttk.Label(left_frame, text="Quantity:").pack(anchor=W)
        self.quantity_entry = ttk.Entry(left_frame, width=30)
        self.quantity_entry.pack(pady=(0, 10), anchor=W)
        
        # Center Frame - Treeview
        center_frame = ttk.Frame(root, padding="10")
        center_frame.pack(side=LEFT, fill=BOTH, expand=True, padx=10, pady=10)
        
        # Search Frame
        search_frame = ttk.Frame(center_frame)
        search_frame.pack(fill=X, pady=(0, 10))
        
        # Product Name Search
        product_search_frame = ttk.Frame(search_frame)
        product_search_frame.pack(side=LEFT, fill=X, expand=True, padx=(0, 5))
        ttk.Label(product_search_frame, text="Search by Product Name:").pack(anchor=W)
        self.product_search_var = tk.StringVar()
        self.product_search_var.trace_add('write', self.on_search_change)
        self.product_search_entry = ttk.Entry(product_search_frame, textvariable=self.product_search_var)
        self.product_search_entry.pack(fill=X)
        
        # Registry Number Search
        registry_search_frame = ttk.Frame(search_frame)
        registry_search_frame.pack(side=LEFT, fill=X, expand=True, padx=(5, 0))
        ttk.Label(registry_search_frame, text="Search by Registry Number:").pack(anchor=W)
        self.registry_search_var = tk.StringVar()
        self.registry_search_var.trace_add('write', self.on_search_change)
        self.registry_search_entry = ttk.Entry(registry_search_frame, textvariable=self.registry_search_var)
        self.registry_search_entry.pack(fill=X)
        
        # Treeview
        self.tree = ttk.Treeview(center_frame, columns=("Product", "Registry", "Quantity"), show="headings")
        self.tree.heading("Product", text="Product Name")
        self.tree.heading("Registry", text="Registry Number")
        self.tree.heading("Quantity", text="Quantity")
        self.tree.pack(fill=BOTH, expand=True)
        
        # Right Frame
        right_frame = ttk.Frame(root, padding="10")
        right_frame.pack(side=RIGHT, fill=Y, padx=10, pady=10)
        
        # Destination Lab Selection
        ttk.Label(right_frame, text="Destination Lab:").pack(anchor=W)
        self.dest_lab_combo = ttk.Combobox(right_frame, width=25, values=list(self.inventory.keys()))
        self.dest_lab_combo.pack(pady=(0, 10), anchor=W)
        self.dest_lab_combo.set("Select Lab")  # Default text
        
        # Bottom Frame - Export Buttons
        bottom_frame = ttk.Frame(root, padding="10")
        bottom_frame.pack(side=BOTTOM, fill=X, padx=10, pady=10)
        
        # Initialize buttons with role-based access
        self.initialize_buttons(bottom_frame, left_frame)

    def create_user_info_label(self):
        """Create label showing logged-in user info"""
        user_frame = ttk.Frame(self.root)
        user_frame.pack(side=TOP, fill=X, padx=10, pady=5)
        
        user_label = ttk.Label(
            user_frame,
            text=f"Logged in as: {self.user_info['username']} ({self.user_info['role']})",
            font=('Helvetica', 10)
        )
        user_label.pack(side=RIGHT)

    def initialize_buttons(self, bottom_frame, left_frame):
        """Initialize buttons with role-based access control"""
        is_admin = self.user_info["role"] == "admin"
        
        # Left frame buttons
        self.add_button = ttk.Button(left_frame, text="Add", command=self.add_product)
        self.add_button.pack(pady=10, anchor=W)
        
        self.edit_button = ttk.Button(left_frame, text="Edit", command=self.edit_product)
        self.edit_button.pack(pady=5, anchor=W)
        if not is_admin:
            self.edit_button["state"] = "disabled"
        
        self.delete_button = ttk.Button(left_frame, text="Delete", command=self.delete_product)
        self.delete_button.pack(pady=5, anchor=W)
        if not is_admin:
            self.delete_button["state"] = "disabled"
        
        # Bottom frame buttons
        self.pdf_button = ttk.Button(bottom_frame, text="PDF", command=self.export_to_pdf)
        self.word_button = ttk.Button(bottom_frame, text="Word", command=self.export_to_word)
        self.excel_button = ttk.Button(bottom_frame, text="Excel", command=self.export_to_excel)
        self.transfer_button = ttk.Button(bottom_frame, text="Transfer", command=self.transfer_product)
        
        # Pack buttons from right to left
        self.transfer_button.pack(side=RIGHT, padx=5)
        self.excel_button.pack(side=RIGHT, padx=5)
        self.word_button.pack(side=RIGHT, padx=5)
        self.pdf_button.pack(side=RIGHT, padx=5)
        
        # Disable export and transfer buttons for non-admin users
        if not is_admin:
            self.transfer_button["state"] = "disabled"
            self.pdf_button["state"] = "disabled"
            self.word_button["state"] = "disabled"
            self.excel_button["state"] = "disabled"

    def on_search_change(self, *args):
        """Handle changes in search fields"""
        self.refresh_treeview()

    def refresh_treeview(self):
        # Clear existing items
        for item in self.tree.get_children():
            self.tree.delete(item)
            
        # Get selected lab
        selected_lab = self.lab_combo.get()
        if selected_lab == "Select Lab":
            return
            
        # Get search terms (convert to lowercase for case-insensitive search)
        product_search = self.product_search_var.get().lower()
        registry_search = self.registry_search_var.get().lower()
            
        # Populate with filtered inventory
        for product in self.inventory[selected_lab]:
            # Check if product matches search criteria
            product_name = product["name"].lower()
            registry = product["registry"].lower()
            
            # Only show items that match both search criteria (if provided)
            if (not product_search or product_search in product_name) and \
               (not registry_search or registry_search in registry):
                self.tree.insert("", "end", values=(
                    product["name"],
                    product["registry"],
                    product["quantity"]
                ))

    def on_lab_selected(self, event=None):
        # Clear search fields when changing labs
        self.product_search_var.set("")
        self.registry_search_var.set("")
        self.refresh_treeview()

    def load_inventory(self):
        """Load inventory from JSON file or return default if file doesn't exist"""
        default_inventory = {
            "Lab 1": [],
            "Lab 2": [],
            "Lab 3": []
        }
        
        try:
            if os.path.exists(self.inventory_file):
                with open(self.inventory_file, 'r') as f:
                    loaded_inventory = json.load(f)
                print("Inventory loaded successfully from file")
                return loaded_inventory
        except json.JSONDecodeError as e:
            print(f"Error reading inventory file: {e}")
            messagebox.showwarning("Warning", "Error reading inventory file. Starting with empty inventory.")
        except Exception as e:
            print(f"Unexpected error loading inventory: {e}")
            messagebox.showwarning("Warning", "Could not load inventory file. Starting with empty inventory.")
        
        return default_inventory

    def save_inventory(self):
        """Save current inventory to JSON file"""
        try:
            with open(self.inventory_file, 'w') as f:
                json.dump(self.inventory, f, indent=4)
            print("Inventory saved successfully")
        except Exception as e:
            print(f"Error saving inventory: {e}")
            messagebox.showerror("Error", "Failed to save inventory to file")

    def load_transfer_log(self):
        """Load transfer log from JSON file or return empty list if file doesn't exist"""
        try:
            if os.path.exists(self.transfer_log_file):
                with open(self.transfer_log_file, 'r') as f:
                    loaded_log = json.load(f)
                print("Transfer log loaded successfully from file")
                return loaded_log
        except json.JSONDecodeError as e:
            print(f"Error reading transfer log file: {e}")
        except Exception as e:
            print(f"Unexpected error loading transfer log: {e}")
        
        return []

    def save_transfer_log(self):
        """Save transfer log to JSON file"""
        try:
            with open(self.transfer_log_file, 'w') as f:
                json.dump(self.transfer_log, f, indent=4)
            print("Transfer log saved successfully")
        except Exception as e:
            print(f"Error saving transfer log: {e}")
            messagebox.showerror("Error", "Failed to save transfer log to file")

    def load_user_log(self):
        """Load user action log from JSON file or return empty list if file doesn't exist"""
        try:
            if os.path.exists(self.user_log_file):
                with open(self.user_log_file, 'r') as f:
                    return json.load(f)
        except Exception as e:
            print(f"Error loading user log: {e}")
        return []

    def save_user_log(self):
        """Save user action log to JSON file"""
        try:
            with open(self.user_log_file, 'w') as f:
                json.dump(self.user_log, f, indent=4)
            print("User log saved successfully")
        except Exception as e:
            print(f"Error saving user log: {e}")

    def log_user_action(self, action, registry, lab, quantity):
        """Log user action with timestamp"""
        log_entry = {
            "user": self.user_info["username"],
            "action": action,
            "registry": registry,
            "lab": lab,
            "quantity": quantity,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        self.user_log.append(log_entry)
        self.save_user_log()
        print(f"Logged user action: {log_entry}")

    def add_product(self):
        # Get values from input fields
        selected_lab = self.lab_combo.get()
        product_name = self.product_entry.get().strip()
        registry_number = self.registry_entry.get().strip()
        quantity = self.quantity_entry.get().strip()
        
        # Validate inputs
        if selected_lab == "Select Lab" or not all([product_name, registry_number, quantity]):
            print("Error: All fields must be filled")
            messagebox.showerror("Error", "All fields must be filled")
            return
        
        try:
            quantity = int(quantity)
            if quantity <= 0:
                raise ValueError("Quantity must be positive")
        except ValueError as e:
            print(f"Error: Invalid quantity - {str(e)}")
            messagebox.showerror("Error", "Quantity must be a positive number")
            return
        
        # Check if product with same registry exists
        existing_product = None
        for product in self.inventory[selected_lab]:
            if product["registry"] == registry_number:
                existing_product = product
                break
        
        if existing_product:
            # Update existing product quantity
            existing_product["quantity"] += quantity
            print(f"Updated quantity for product {registry_number} in {selected_lab}")
        else:
            # Create new product
            new_product = {
                "name": product_name,
                "registry": registry_number,
                "quantity": quantity
            }
            self.inventory[selected_lab].append(new_product)
            print(f"Added new product to {selected_lab}: {new_product}")
        
        # Log the action
        self.log_user_action("add", registry_number, selected_lab, quantity)
        
        # Save inventory to file
        self.save_inventory()
        
        # Refresh display and clear inputs
        self.refresh_treeview()
        self.product_entry.delete(0, END)
        self.registry_entry.delete(0, END)
        self.quantity_entry.delete(0, END)

    def transfer_product(self):
        # Get source and destination labs
        from_lab = self.lab_combo.get()
        to_lab = self.dest_lab_combo.get()
        
        # Validate lab selection
        if from_lab == "Select Lab" or to_lab == "Select Lab":
            messagebox.showerror("Error", "Please select both source and destination labs")
            return
        
        if from_lab == to_lab:
            messagebox.showerror("Error", "Source and destination labs must be different")
            return
        
        # Get selected item from treeview
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showerror("Error", "Please select a product to transfer")
            return
        
        # Get product details from selected item
        item = self.tree.item(selected_items[0])
        product_name, registry, quantity = item['values']
        
        # Ask for transfer quantity
        transfer_quantity = self.ask_transfer_quantity(quantity)
        if not transfer_quantity:
            return
        
        # Find and update product in source lab
        source_product = None
        for product in self.inventory[from_lab]:
            if product["registry"] == registry:
                source_product = product
                break
        
        if not source_product:
            messagebox.showerror("Error", "Product not found in source lab")
            return
        
        # Update source lab quantity
        source_product["quantity"] -= transfer_quantity
        
        # Remove product if quantity becomes 0
        if source_product["quantity"] == 0:
            self.inventory[from_lab].remove(source_product)
        
        # Find or create product in destination lab
        dest_product = None
        for product in self.inventory[to_lab]:
            if product["registry"] == registry:
                dest_product = product
                break
        
        if dest_product:
            dest_product["quantity"] += transfer_quantity
        else:
            self.inventory[to_lab].append({
                "name": product_name,
                "registry": registry,
                "quantity": transfer_quantity
            })
        
        # Create transfer log entry
        transfer_entry = {
            "product_name": product_name,
            "registry": registry,
            "quantity": transfer_quantity,
            "from_lab": from_lab,
            "to_lab": to_lab,
            "transfer_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        self.transfer_log.append(transfer_entry)
        
        # Log the user action
        self.log_user_action("transfer", registry, f"{from_lab} -> {to_lab}", transfer_quantity)
        
        # Save both inventory and transfer log
        self.save_inventory()
        self.save_transfer_log()
        
        # Refresh display
        self.refresh_treeview()
        messagebox.showinfo("Success", f"Transferred {transfer_quantity} units of {product_name} from {from_lab} to {to_lab}")

    def ask_transfer_quantity(self, max_quantity):
        """Ask user for transfer quantity using a dialog"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Transfer Quantity")
        dialog.geometry("300x150")
        dialog.transient(self.root)
        dialog.grab_set()
        
        # Center the dialog
        dialog.geometry("+%d+%d" % (
            self.root.winfo_rootx() + self.root.winfo_width()/2 - 150,
            self.root.winfo_rooty() + self.root.winfo_height()/2 - 75
        ))
        
        # Add widgets
        ttk.Label(dialog, text=f"Enter quantity to transfer (max {max_quantity}):").pack(pady=10)
        quantity_var = tk.StringVar()
        entry = ttk.Entry(dialog, textvariable=quantity_var)
        entry.pack(pady=5)
        
        result = [None]  # Use list to store result
        
        def validate_and_close():
            try:
                quantity = int(quantity_var.get())
                if 0 < quantity <= int(max_quantity):
                    result[0] = quantity
                    dialog.destroy()
                else:
                    messagebox.showerror("Error", f"Please enter a number between 1 and {max_quantity}")
            except ValueError:
                messagebox.showerror("Error", "Please enter a valid number")
        
        ttk.Button(dialog, text="OK", command=validate_and_close).pack(pady=10)
        
        # Wait for dialog to close
        dialog.wait_window()
        return result[0]

    def create_export_directories(self):
        """Create export directories if they don't exist"""
        for directory in self.export_dirs.values():
            os.makedirs(directory, exist_ok=True)
            print(f"Ensured export directory exists: {directory}")

    def export_to_excel(self):
        """Export current lab's inventory to Excel file"""
        # Get selected lab
        selected_lab = self.lab_combo.get()
        
        # Validate selection
        if selected_lab == "Select Lab":
            messagebox.showerror("Error", "Please select a lab first")
            return
            
        # Get inventory for selected lab
        lab_inventory = self.inventory[selected_lab]
        
        # Check if inventory is empty
        if not lab_inventory:
            messagebox.showerror("Error", f"No inventory items found in {selected_lab}")
            return
            
        try:
            # Convert inventory to pandas DataFrame
            df = pd.DataFrame(lab_inventory)
            
            # Rename columns to match desired format
            df = df.rename(columns={
                'name': 'Product Name',
                'registry': 'Registry Number',
                'quantity': 'Quantity'
            })
            
            # Reorder columns
            df = df[['Product Name', 'Registry Number', 'Quantity']]
            
            # Generate filename in excel directory
            filename = os.path.join(
                self.export_dirs['excel'],
                f"{selected_lab.replace(' ', '')}_inventory.xlsx"
            )
            
            # Export to Excel
            df.to_excel(filename, index=False, engine='openpyxl')
            
            messagebox.showinfo("Success", f"Inventory exported successfully to {filename}")
            print(f"Exported inventory to {filename}")
            
        except Exception as e:
            error_msg = f"Error exporting to Excel: {str(e)}"
            print(error_msg)
            messagebox.showerror("Error", error_msg)

    def export_to_pdf(self):
        """Export current lab's inventory to PDF file"""
        # Get selected lab
        selected_lab = self.lab_combo.get()
        
        # Validate selection
        if selected_lab == "Select Lab":
            messagebox.showerror("Error", "Please select a lab first")
            return
            
        # Get inventory for selected lab
        lab_inventory = self.inventory[selected_lab]
        
        # Check if inventory is empty
        if not lab_inventory:
            messagebox.showerror("Error", f"No inventory items found in {selected_lab}")
            return
            
        try:
            # Generate filename in pdf directory
            filename = os.path.join(
                self.export_dirs['pdf'],
                f"{selected_lab.replace(' ', '')}_inventory.pdf"
            )
            
            # Create PDF document
            doc = SimpleDocTemplate(
                filename,
                pagesize=letter,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Create the table data
            elements = []
            
            # Add title
            title = f"Inventory Report for {selected_lab}"
            
            # Define table headers and data
            headers = ["Product Name", "Registry Number", "Quantity"]
            data = [headers]  # Start with headers
            
            # Add inventory data
            for item in lab_inventory:
                data.append([
                    item["name"],
                    item["registry"],
                    str(item["quantity"])
                ])
            
            # Create table
            table = Table(data)
            
            # Add style to table
            style = TableStyle([
                # Title style
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),  # Center align header row
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),  # Grey background for header
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),  # White text for header
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Bold font for header
                ('FONTSIZE', (0, 0), (-1, 0), 12),  # Header font size
                
                # Table body style
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),  # Left align body text
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),  # Regular font for body
                ('FONTSIZE', (0, 1), (-1, -1), 10),  # Body font size
                
                # Grid style
                ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Add grid lines
                ('BOX', (0, 0), (-1, -1), 2, colors.black),  # Add outer border
                
                # Alternating row colors
                ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.lightgrey])
            ])
            
            table.setStyle(style)
            
            # Create title
            title_style = TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 16),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
            ])
            
            title_table = Table([[title]], colWidths=[400])
            title_table.setStyle(title_style)
            
            # Add elements to document
            elements.append(title_table)
            elements.append(table)
            
            # Build PDF
            doc.build(elements)
            
            messagebox.showinfo("Success", f"Inventory exported successfully to {filename}")
            print(f"Exported inventory to {filename}")
            
        except Exception as e:
            error_msg = f"Error exporting to PDF: {str(e)}"
            print(error_msg)
            messagebox.showerror("Error", error_msg)

    def export_to_word(self):
        """Export current lab's inventory to Word document"""
        # Get selected lab
        selected_lab = self.lab_combo.get()
        
        # Validate selection
        if selected_lab == "Select Lab":
            messagebox.showerror("Error", "Please select a lab first")
            return
            
        # Get inventory for selected lab
        lab_inventory = self.inventory[selected_lab]
        
        # Check if inventory is empty
        if not lab_inventory:
            messagebox.showerror("Error", f"No inventory items found in {selected_lab}")
            return
            
        try:
            # Create Word document
            doc = Document()
            
            # Add title
            title = f"Inventory Report for {selected_lab}"
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add some space after title
            doc.add_paragraph()
            
            # Create table
            table = doc.add_table(rows=1, cols=3, style='Table Grid')
            table.autofit = True
            
            # Set header row
            header_cells = table.rows[0].cells
            headers = ["Product Name", "Registry Number", "Quantity"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Make headers bold
                run = header_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
            
            # Add inventory data
            for item in lab_inventory:
                row_cells = table.add_row().cells
                row_cells[0].text = item["name"]
                row_cells[1].text = item["registry"]
                row_cells[2].text = str(item["quantity"])
                
                # Center align quantity
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set column widths
            for column in table.columns:
                for cell in column.cells:
                    cell.width = Inches(2.0)
                    # Vertically center align all cells
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add timestamp at the bottom
            doc.add_paragraph()
            timestamp = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Generate filename in word directory
            filename = os.path.join(
                self.export_dirs['word'],
                f"{selected_lab.replace(' ', '')}_inventory.docx"
            )
            doc.save(filename)
            
            messagebox.showinfo("Success", f"Inventory exported successfully to {filename}")
            print(f"Exported inventory to {filename}")
            
        except Exception as e:
            error_msg = f"Error exporting to Word: {str(e)}"
            print(error_msg)
            messagebox.showerror("Error", error_msg)

    def delete_product(self):
        """Delete or reduce quantity of selected product"""
        # Get selected lab
        selected_lab = self.lab_combo.get()
        
        # Validate lab selection
        if selected_lab == "Select Lab":
            messagebox.showerror("Error", "Please select a lab first")
            return
        
        # Get selected item from treeview
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showerror("Error", "Please select a product to delete")
            return
        
        # Get product details from selected item
        item = self.tree.item(selected_items[0])
        product_name, registry, current_quantity = item['values']
        current_quantity = int(current_quantity)
        
        # Find product in inventory
        product_to_delete = None
        for product in self.inventory[selected_lab]:
            if product["registry"] == registry:
                product_to_delete = product
                break
        
        if not product_to_delete:
            messagebox.showerror("Error", "Product not found in inventory")
            return
        
        # Ask for quantity to delete
        delete_quantity = simpledialog.askinteger(
            "Delete Product",
            f"How many units do you want to delete?\nCurrent quantity: {current_quantity}",
            parent=self.root,
            minvalue=1,
            maxvalue=current_quantity
        )
        
        if delete_quantity is None:  # User cancelled
            return
        
        try:
            # Update or remove product
            if delete_quantity == current_quantity:
                # Remove entire product
                self.inventory[selected_lab].remove(product_to_delete)
                messagebox.showinfo("Success", f"Removed {product_name} from inventory")
            else:
                # Update quantity
                product_to_delete["quantity"] -= delete_quantity
                messagebox.showinfo("Success", 
                    f"Deleted {delete_quantity} units of {product_name}\n"
                    f"New quantity: {product_to_delete['quantity']}")
            
            # Log the action
            self.log_user_action("delete", registry, selected_lab, delete_quantity)
            
            # Save inventory and refresh display
            self.save_inventory()
            self.refresh_treeview()
            
        except Exception as e:
            error_msg = f"Error deleting product: {str(e)}"
            print(error_msg)
            messagebox.showerror("Error", error_msg)

    def edit_product(self):
        """Edit selected product's name and quantity"""
        # Get selected lab
        selected_lab = self.lab_combo.get()
        
        # Validate lab selection
        if selected_lab == "Select Lab":
            messagebox.showerror("Error", "Please select a lab first")
            return
        
        # Get selected item from treeview
        selected_items = self.tree.selection()
        if not selected_items:
            messagebox.showerror("Error", "Please select a product to edit")
            return
        
        # Get product details from selected item
        item = self.tree.item(selected_items[0])
        product_name, registry, current_quantity = item['values']
        current_quantity = int(current_quantity)
        
        # Find product in inventory
        product_to_edit = None
        for product in self.inventory[selected_lab]:
            if product["registry"] == registry:
                product_to_edit = product
                break
        
        if not product_to_edit:
            messagebox.showerror("Error", "Product not found in inventory")
            return
        
        # Open edit dialog
        dialog = EditDialog(
            self.root,
            "Edit Product",
            product_name,
            current_quantity
        )
        
        # Wait for dialog to close
        self.root.wait_window(dialog.dialog)
        
        # Check if user saved changes
        if dialog.result:
            try:
                # Calculate quantity change
                quantity_change = dialog.result["quantity"] - current_quantity
                
                # Update product
                product_to_edit["name"] = dialog.result["name"]
                product_to_edit["quantity"] = dialog.result["quantity"]
                
                # Log the action
                self.log_user_action("edit", registry, selected_lab, quantity_change)
                
                # Save inventory and refresh display
                self.save_inventory()
                self.refresh_treeview()
                
                messagebox.showinfo("Success", 
                    f"Updated product details:\n"
                    f"Name: {dialog.result['name']}\n"
                    f"Quantity: {dialog.result['quantity']}")
                
            except Exception as e:
                error_msg = f"Error updating product: {str(e)}"
                print(error_msg)
                messagebox.showerror("Error", error_msg)

    def show_logs(self):
        """Open the log viewer window"""
        LogViewer(self.root)

def main():
    # Show login window first
    login = LoginWindow()
    authenticated, user_info = login.run()
    
    if authenticated:  # Only proceed if authentication successful
        # Create and show main application window
        root = tk.Tk()
        app = LabInventoryManager(root, user_info)
        root.mainloop()

if __name__ == "__main__":
    main() 